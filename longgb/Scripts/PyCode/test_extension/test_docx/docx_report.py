# -*- coding:utf-8 -*-
# @Author  : 'longguangbin'
# @Contact : lgb453476610@163.com
# @Date    : 2018/12/27
"""  
Usage Of 'docx_report' : 
"""

import os
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor


class ReportWord:
    def __init__(self, report_name, period, file_path, data, sim_screen, total_sku, screen_sku):
        # report_name = '11977'
        # period = ['2016-07-01', '2016-10-24']
        # file_path = r'D:\Lgb\ReadData\RealityAnalysis\report'
        self.document = Document()
        self.report_name = report_name
        self.period = period
        self.file_path = file_path
        self.data = data
        self.sim_screen = sim_screen
        self.total_sku = total_sku
        self.screen_sku = screen_sku

    def pinleianalysis(self, pinlei, pinleiid):
        '''
        A.品类分析
        :return:
        '''
        # pinlei = [u'厨具',u'厨房配件',u'储物/置物架']; pinleiid = 11977; skusum= total_sku
        self.document.add_heading(u'A 品类概况', 0)
        self.document.add_heading(u'a) 数据准备', 1)
        text1 = u'''\t数据选自“{0}”-“{1}”下的“{2}”（三级分类：{3}）共{4}个SKU（以2016-11-19日上柜的北京RDC、自营SKU为准）。为满足仿真条件，在此基础上，我们按照以下标准对SKU进行了筛选。'''
        self.tempnum = self.data.item_sku_id.unique().size - self.sim_screen.three_conditions.sum()
        self.document.add_paragraph(text1.format(pinlei[0], pinlei[1], pinlei[2], pinleiid, self.total_sku))
        self.document.add_paragraph(u"\tSKU总数为{0}个，在分析首日之前上柜的SKU总数为{1}个：".format(self.total_sku, self.screen_sku))
        self.document.add_paragraph(u"\t1）首日销量预测为空的sku数量为：%d" % self.sim_screen.first_day_ofdsales.sum())
        self.document.add_paragraph(u"\t2）首日库存数据为0的sku数量为：%d" % self.sim_screen.first_day_inv.sum())
        self.document.add_paragraph(u"\t3）期间总销量为0的sku数量为：%d" % self.sim_screen.total_sales.sum())
        self.document.add_paragraph(u"\t以上情况至少出现一次的sku数量为{0}个，满足筛选条件sku个数为{1}个（占总数{2:.2f}%）用于仿真。".format(
            self.sim_screen.three_conditions.sum(), self.tempnum, self.tempnum / self.total_sku * 100))
        runname = self.document.add_paragraph().add_run(
            u'\t注. 1）为展现该三级品类库存基本情况，采用{0}个SKU进行分析；2）为满足仿真条件，仅采用958个SKU进行仿真对比。'.format(self.screen_sku, self.tempnum))
        fontname = runname.font
        fontname.italic = True
        fontname.color.rgb = RGBColor(54, 95, 145)
        self.count_supp_num = self.data[self.data.supp_brevity_cd.isnull() == False][
            "supp_brevity_cd"].drop_duplicates().count()
        self.count_caigou = self.data[self.data.pur_bill_id.isnull() == False]["pur_bill_id"].count()
        self.document.add_heading(u'b) 品类描述性统计', 1)
        self.document.add_paragraph(u'\t分析期间的描述性统计分析：')
        self.document.add_paragraph(u'\t1. SKU总数为：{0}个'.format(self.screen_sku))
        self.document.add_paragraph(u'\t2. 分析期间总采购次数为：{0}次'.format(self.count_caigou))
        self.document.add_paragraph(u'\t3. 分析期间供应商总数为：{0}个'.format(self.count_supp_num))
        self.document.add_picture(self.file_path + os.sep + 'caigou.png', height=Inches(3.84), width=Inches(5.76))

    def skuget(self, name, title, word1, word2):
        name_path = self.file_path + os.sep + name
        self.document.add_heading(title, 2)
        self.document.add_paragraph(word1)
        # high_cr_low_td_path = file_path + os.sep + u'high_cr_low_td'
        for each in os.listdir(name_path)[:2]:
            self.document.add_paragraph(u'\tSKU:{0}'.format(each[:-4]))
            self.document.add_picture(name_path + os.sep + each, height=Inches(2.88), width=Inches(5.76))
        self.document.add_paragraph(word2)

    def kpianalysis(self, kpi_cr_str, kpi_ito_str, sim_date_range):
        self.document.add_heading(u'B KPI分析', 0)
        self.document.add_paragraph(u'\t目前库存系统的主要KPI 是现货率和库存周转天数。下面我们对“储物/置物架”类SKU 的这两个指标进行分析。')
        self.document.add_paragraph(
            u'\t注：KPI的分析当中会用到中位数。中位数顾名思义，处于中间位置的数，其可将数值集合划分为相等的上下两部分，中位数不会受到少量异常值的影响，而如果存在异常值，均值的变化会比较明显，受异常值影响较大。')
        self.document.add_heading(u'a) KPI计算口径', 1)
        table = self.document.add_table(1, 3)
        # table.style = 'Light Shading Accent 1'
        # 标题
        heading_cells = table.rows[0].cells
        heading_cells[0].text = u'颗粒度'
        heading_cells[1].text = u'现货率(CR)'
        heading_cells[2].text = u'周转(ITO)'
        # 添加内容
        items = [[u'SKU', u'可售天数/总天数', u'平均现货库存/日均销量'], [u'三级品类', u'可售SKU数/上柜SKU总数\n可售SKU数：上柜有货或上柜可预订', u'平均现货库存/日均销量']]
        for item in items:
            cells = table.add_row().cells
            cells[0].text = item[0]
            cells[1].text = item[1]
            cells[2].text = item[2]
        runname = self.document.add_paragraph().add_run(
            u'\t注. 1)以上平均值期间为{0}至{1}。'.format(sim_date_range[0], sim_date_range[1]))
        fontname = runname.font
        fontname.italic = True
        fontname.color.rgb = RGBColor(54, 95, 145)
        self.document.add_heading(u'b) 现货率', 1)
        self.document.add_paragraph(u'\t下图展示了所有SKU 的现货率频数分布。')
        self.document.add_picture(self.file_path + os.sep + 'cr2.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\t' + kpi_cr_str)
        self.document.add_heading(u'c) 周转天数', 1)
        self.document.add_paragraph(u'\t下图展示了所有SKU 的周转天数频数分布。')
        self.document.add_picture(self.file_path + os.sep + 'td2.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\t' + kpi_ito_str)
        self.document.add_heading(u'd) 现货率、周转天数联合分析', 1)
        self.document.add_paragraph(u'\t根据每个SKU周转天数（横轴）和现货率（纵轴）所作出的散点图。如下图所示：')
        self.document.add_picture(self.file_path + os.sep + 'cr_td.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\t图中横轴是0.8，纵轴是60。从图中，我们可以看出，库存周转天数跟现货率呈现出明显的正相关性，即库存周转率越高，现货率越低。')
        self.document.add_paragraph(u'\t根据现货率的高低和库存周转天数的大小，我们将图中分为4个区域：')
        self.document.add_paragraph(u'\t（1）现货率高，周转低的SKU')
        self.document.add_paragraph(u'\t（2）现货率高，周转高的SKU')
        self.document.add_paragraph(u'\t（3）现货率低，周转低的SKU')
        self.document.add_paragraph(u'\t（4）现货率低，周转高的SKU')
        self.document.add_paragraph(u'\t主要优化空间在（2）、（3）和（4）。（2）区域适当降低周转，（3）区域适当提高现货率，（4）区域周转和现货率都有优化空间。')
        self.document.add_heading(u'e) 现货率、周转天数case study', 1)
        self.document.add_paragraph(u'\t实例说明：下面对这四个区域的SKU逐一进行分析。')
        self.skuget(u'high_cr_low_td', u'（1）现货率高，周转低的SKU', u'\t这部分SKU以较低的周转实现了较高的现货率。下面给出几个此类SKU代表。',
                    u'\t这部分SKU以较低的周转实现了较高的现货率。')
        temp_text2 = u'a'
        self.skuget(u'high_cr_high_td', u'（2）现货率高，周转高的SKU',
                    u'\t这部分SKU存在降低库存周转的空间，尤其在现货率接近1的SKU。下图展示了几个比较典型的SKU的库存情况。下面给出几个此类SKU代表。',
                    u'\t在这个区域的SKU，库存水平（相对销量）始终保持在较高的水平。由库存水平可以看出，周转高的主要原因是：（a）单次采购量过大。（b）采购时间过早。\n\t对于这部分SKU，可以采用降低单次采购量，保持高服务水平的情况下多次采购来降低周转。如果适当推迟补货时间，现货率仍然会维持在100%（或接近100%），同时库存周转天数可以大幅降低。')
        self.skuget(u'low_cr_low_td', u'（3）现货率低，周转低的SKU',
                    u'\t在同样库存周转天数的情况下，这部分SKU的现货率明显低于位于（1）的SKU（现货率高，周转低的SKU）。下面给出几个此类SKU代表。',
                    u'\t导致这类SKU现货率较低/周转天数长的原因为：\n\t（a）采购时间把握不准，采购时间晚（存在库存将至0才补货的现象）；\n\t（b）存在长时间缺货现象，可能的原因是采销没有及时补货或者补货了但是供应商没有及时送到。')
        self.skuget(u'low_cr_high_td', u'（4）现货率低，周转高的SKU',
                    u'\t这部分SKU的现货率和库存周转天数都存在改进的空间，其表现不佳一般是上述各原因综合导致的。下面给出几个此类SKU代表。',
                    u'\t综上所述，我们认为导致现货率低或者周转天数长的主要原因有：\n\t（a）采购时间过早/过晚；\n\t（b）单次订货量过大；\n\t（c）供应商供货不稳定。')

    def zcaseget(self):
        name_path = self.file_path + os.sep + u'zcase'
        for each in os.listdir(name_path)[:3]:
            self.document.add_paragraph(u'\tSKU:{0}'.format(each[:-4]))
            self.document.add_picture(name_path + os.sep + each, height=Inches(2.88), width=Inches(5.76))

    def buhuoanalysis(self):
        self.document.add_heading(u'C 补货时间分析', 0)
        self.document.add_heading(u'a) 评判指标Z引入', 1)
        self.document.add_paragraph(u'\t在这部分，我们设计量化指标评价采购下单时间。对于每一个采购单，我们用s表示补货点，Q表示补货量（即s+Q是目标库存）。我们按照如下公式定义Z指标：')
        self.document.add_paragraph(
            u'\t其中，d为VLT期间日均销量预测（报告中的结果使用的是历史28天平均销量；我们也尝试了使用未来7/14/28天实际平均销量，订单当天对未来销量的预测值，得到的结果一致），σ为日销量的标准差，L为VLT天数（假设采销对于每一个采购单的VLT有较准确的估计；如果VLT波动过大，可以略微修改Z指标来包括随机VLT的影响）。')
        self.document.add_heading(u'b) 评判指标Z值分析', 1)
        self.document.add_paragraph(u'\t共有{0}个有效采购单，对于每一个采购单/SKU组合计算Z指标，其分布如下：'.format(self.count_caigou))
        self.document.add_picture(self.file_path + os.sep + 'z.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\tZ指标能够较好的衡量补货时间点的早晚。对于一次补货来讲，可以用这批货到货的前一天是否缺货来衡量补货时间的早晚。')
        self.document.add_heading(u'c) 评判指标Z case study', 1)
        self.document.add_paragraph(u'\t实例说明。')
        self.zcaseget()

    def bpanalysis(self):
        self.document.add_heading(u'D 补货量（BP）分析', 0)
        self.document.add_picture(self.file_path + os.sep + 'bp2.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_heading(u'A-D band的BP值分析', 1)
        self.document.add_picture(self.file_path + os.sep + 'bp3.png', height=Inches(3.84), width=Inches(5.76))

    def suppanalysis(self):
        self.document.add_heading(u'E 供应商分析', 0)
        self.document.add_paragraph(
            u'\t供应商的稳定性是影响库存情况的一个重要因素。下面从供应商的vlt、满足率两个方面来分析。总共有{0}个供应商。'.format(self.count_supp_num))
        self.document.add_paragraph(
            u'\t我们为每个供应商的表现（采购单量，VLT平均天数，标准差，变异系数，订单满足率，订单满足率标准差）从0到10打分，结果如下表所示。(在任一指标上，表现最好的供应商为10分，最差的供应商为1分，其余供应商分数根据线性插值确定。)')
        self.document.add_paragraph(
            u'\t注：\n\t1.VLT（Vendor lead time）：供应商送货提前期\n\t2.订单满足率口径：实际满足率=实际到货量/原始采购量\n\t3.变异系数主要应用在评价数据的波动性，避免量纲带来的影响')
        self.document.add_heading(u'a) VLT变异系数', 1)
        self.document.add_paragraph(u'\t下图为供应商维度vlt的变异系数(CV)分布：')
        self.document.add_picture(self.file_path + os.sep + 'vltcv.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_heading(u'b) 满足率', 1)
        self.document.add_paragraph(u'\t供应商维度满足率计算口径：实际满足率=实际到货量/原始采购量\n\t下面两图为供应商两种口径的满足率分布。分别为实际满足率和实际相对回告满足率')
        self.document.add_picture(self.file_path + os.sep + 'manzu.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_picture(self.file_path + os.sep + 'manzu2.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_heading(u'c) VLT变异系数、满足率联合分析', 1)
        self.document.add_picture(self.file_path + os.sep + 'vltcv_manzu.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\t区域（1）：供应商波动性小，满足率高。供应商在波动性和满足率上表现最好的。')
        self.document.add_paragraph(u'\t区域（2）：供应商波动性大，满足率高。为避免供应商波动性带来的服务水平下降，提早采购则有周转变高的风险。')
        self.document.add_paragraph(u'\t区域（3）：供应商波动性小，满足率小。供应商在波动性上表现很好，满足率很低，容易造成服务水平降低。')
        self.document.add_paragraph(u'\t区域（4）：供应商波动性大，满足率低。供应商在波动性和满足率上表现最差的，极易造成低服务水平，同时如果为了避免供应商波动性，提早采购则有造成高周转的风险。')

    def simanalysis(self):
        self.document.add_heading(u'F 模拟仿真', 0)
        self.document.add_paragraph(
            u'\t最终同时满足三条仿真条件的sku共计{0}个。占满足分析条件{1}个sku的{2:.2f}%。'.format(self.tempnum, self.screen_sku,
                                                                        self.tempnum / self.screen_sku * 100))
        self.document.add_picture(self.file_path + os.sep + 'ito_his_sim.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\t图中两个数字分别为在红线上下两部分的sku个数。红线上部分代表周转变大的情况，红线下部分代表周转变小的情况。')
        self.document.add_picture(self.file_path + os.sep + 'cr_his_sim.png', height=Inches(3.84), width=Inches(5.76))
        self.document.add_paragraph(u'\t图中两个数字分别为在红线上下两部分的sku个数。红线上部分代表现货率变大的情况，红线下部分代表现货率变小的情况。')
        self.document.add_paragraph(u'\t综上所述：造成低现货率或者高周转的主要原因有：\n\t1、采购时机不好\n\t2、供应商送货不稳定\n\t3、单次采购量不合理')

    def closeword(self):
        # self.document.add_page_break()
        self.document.save(self.file_path + os.sep + 'report_word_{0}.docx'.format(self.report_name))

    def generateword(self, pinlei, pinleiid, kpi_cr_str, kpi_ito_str):
        self.pinleianalysis(pinlei, pinleiid)
        self.kpianalysis(kpi_cr_str, kpi_ito_str)
        self.buhuoanalysis()
        self.bpanalysis()
        self.suppanalysis()
        self.simanalysis()
        self.closeword()
