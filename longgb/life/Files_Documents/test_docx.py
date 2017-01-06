#-*- coding:utf-8 -*-
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt

document = docx.Document()
table = document.add_table(rows=1, cols=3)
# table.style = 'Light Grid Accent 1'
# table.style = 'Light List Accent 1'
table.style = 'Light List Accent 1'
table.columns[0].width = Inches(0.88)
table.columns[1].width = Inches(2.95)
table.columns[2].width = Inches(1.96)
hdr_cells = table.rows[0].cells
# hdr_cells[0].paragraphs = 'Qty'
hdr_cells[0].text = 'Qty'
# hdr_cells[0].paragraph.size = Pt(10.5)
# hdr_cells[0].paragraphs.style = 'Body Text 2'
# hdr_cells[0].paragraphs.size = Pt(10.5)
# hdr_cells[0].paragraphs.font.size = Pt(10.5)
# print hdr_cells[0].paragraphs
# runname = hdr_cells[0].add_paragraph().add_run('Qty')
# hdr_cells[0].add_paragraph('Qty')
# fontname = runname.font
# fontname.italic = True
# fontname.size = Pt(10.5)
# font_s = hdr_cells[0].font
# font_s.size = Pt(10.5)
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
item = ['Qty', 'Id', 'Desc']
cells = table.add_row().cells
cells[0].text = item[0]
cells[1].text = item[1]
cells[2].text = item[2]
cells = table.add_row().cells
cells[0].text = item[0]
cells[1].text = item[1]
cells[2].text = item[2]
cells = table.add_row().cells
cells[0].text = item[0]
cells[1].text = item[1]
cells[2].text = item[2]
document.save(r'C:\Users\sunwenxiu\Desktop\wenjian\demo.docx')


